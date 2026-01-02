#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理 MCP 服务器

提供 PDF、Word、Markdown、HTML、TXT 等格式的文档处理功能。
使用标准 MCP 协议，可通过 .mcp.json 配置启动。
"""

import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional
from abc import ABC, abstractmethod
import asyncio

# 添加项目根目录到 Python 路径
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# 导入配置
from mcp_servers.config import (
    SUPPORTED_FORMATS,
    ALLOWED_EXTENSIONS,
    MAX_FILE_SIZE,
    MAX_BATCH_SIZE,
    MAX_SUMMARY_LENGTH,
    MAX_CONTENT_LENGTH,
    DEFAULT_ENCODING,
    is_supported_format,
    get_format_type,
    ensure_temp_dir,
)

# 导入标准 MCP 协议
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

# 尝试导入文档处理库
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False


# ============================================================
# 文档处理器基类
# ============================================================

class BaseDocumentProcessor(ABC):
    """文档处理器基类"""

    def __init__(self):
        self.format_name = "unknown"

    @abstractmethod
    def read(self, file_path: str) -> Dict[str, Any]:
        """读取文档内容"""
        pass

    @abstractmethod
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取文档元数据"""
        pass

    def validate_file(self, file_path: str) -> tuple[bool, str]:
        """验证文件"""
        if not os.path.exists(file_path):
            return False, "文件不存在"

        if not os.path.isfile(file_path):
            return False, "路径不是文件"

        if not is_supported_format(file_path):
            return False, f"不支持的文件格式"

        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return False, f"文件过大 ({file_size} > {MAX_FILE_SIZE})"

        return True, ""


# ============================================================
# 具体文档处理器
# ============================================================

class PDFProcessor(BaseDocumentProcessor):
    """PDF 文档处理器"""

    def __init__(self):
        self.format_name = "pdf"

    def read(self, file_path: str) -> Dict[str, Any]:
        """读取 PDF 内容"""
        if not HAS_PYPDF2:
            return {
                "success": False,
                "error": "PyPDF2 未安装，请运行: pip install PyPDF2"
            }

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                content_parts = []

                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)

                content = "\n\n".join(content_parts)

                return {
                    "success": True,
                    "content": content,
                    "page_count": len(pdf_reader.pages),
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"读取 PDF 失败: {str(e)}"
            }

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取 PDF 元数据"""
        if not HAS_PYPDF2:
            return {"error": "PyPDF2 未安装"}

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "format": "PDF",
                }

                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                        "producer": pdf_reader.metadata.get('/Producer', ''),
                    })

                file_stat = os.stat(file_path)
                metadata.update({
                    "file_size": file_stat.st_size,
                    "created": file_stat.st_ctime,
                    "modified": file_stat.st_mtime,
                })

                return metadata
        except Exception as e:
            return {"error": f"获取元数据失败: {str(e)}"}


class WordProcessor(BaseDocumentProcessor):
    """Word 文档处理器"""

    def __init__(self):
        self.format_name = "word"

    def read(self, file_path: str) -> Dict[str, Any]:
        """读取 Word 内容"""
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx 未安装，请运行: pip install python-docx"
            }

        try:
            doc = Document(file_path)
            content_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            content = "\n\n".join(content_parts)

            return {
                "success": True,
                "content": content,
                "paragraph_count": len(doc.paragraphs),
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"读取 Word 文档失败: {str(e)}"
            }

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取 Word 元数据"""
        if not HAS_DOCX:
            return {"error": "python-docx 未安装"}

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            metadata = {
                "format": "Word",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "comments": core_props.comments or "",
            }

            file_stat = os.stat(file_path)
            metadata.update({
                "file_size": file_stat.st_size,
                "created": file_stat.st_ctime,
                "modified": file_stat.st_mtime,
            })

            return metadata
        except Exception as e:
            return {"error": f"获取元数据失败: {str(e)}"}


class MarkdownProcessor(BaseDocumentProcessor):
    """Markdown 文档处理器"""

    def __init__(self):
        self.format_name = "markdown"

    def read(self, file_path: str) -> Dict[str, Any]:
        """读取 Markdown 内容"""
        try:
            # 使用备用编码机制读取
            content = self._read_with_fallback(file_path)

            # 统计信息
            lines = content.split('\n')
            code_blocks = [line for line in lines if line.strip().startswith('```')]

            return {
                "success": True,
                "content": content,
                "line_count": len(lines),
                "code_block_count": len(code_blocks) // 2,
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"读取 Markdown 失败: {str(e)}"
            }

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取 Markdown 元数据"""
        try:
            # 使用备用编码机制读取
            content = self._read_with_fallback(file_path)

            # 获取实际使用的编码
            encoding = self._get_actual_encoding(file_path)

            lines = content.split('\n')
            file_stat = os.stat(file_path)

            # 提取标题
            headings = [line for line in lines if line.strip().startswith('#')]

            metadata = {
                "format": "Markdown",
                "line_count": len(lines),
                "heading_count": len(headings),
                "encoding": encoding,
                "file_size": file_stat.st_size,
                "created": file_stat.st_ctime,
                "modified": file_stat.st_mtime,
            }

            return metadata
        except Exception as e:
            return {"error": f"获取元数据失败: {str(e)}"}

    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        if HAS_CHARDET:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                detected = result['encoding'] or DEFAULT_ENCODING
                # 验证检测的编码是否可用
                try:
                    raw_data.decode(detected)
                    return detected
                except (UnicodeDecodeError, LookupError):
                    pass  # 检测失败，使用默认编码
        return DEFAULT_ENCODING

    def _read_with_fallback(self, file_path: str) -> str:
        """使用备用编码列表读取文件"""
        # 获取检测到的编码
        detected_encoding = self._detect_encoding(file_path)

        # 编码尝试顺序：检测编码 -> 默认编码 -> 备用编码列表
        encodings_to_try = [detected_encoding, DEFAULT_ENCODING] + ['utf-8', 'gbk', 'gb2312', 'latin1']

        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue

        # 如果所有编码都失败，返回错误
        raise ValueError(f"无法解码文件，尝试的编码: {encodings_to_try}")

    def _get_actual_encoding(self, file_path: str) -> str:
        """获取实际可用的编码"""
        detected_encoding = self._detect_encoding(file_path)
        encodings_to_try = [detected_encoding, DEFAULT_ENCODING, 'utf-8', 'gbk', 'gb2312', 'latin1']

        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    f.read(1)  # 尝试读取一个字符
                return encoding
            except (UnicodeDecodeError, LookupError):
                continue

        return DEFAULT_ENCODING


class HTMLProcessor(BaseDocumentProcessor):
    """HTML 文档处理器"""

    def __init__(self):
        self.format_name = "html"

    def read(self, file_path: str) -> Dict[str, Any]:
        """读取 HTML 内容"""
        try:
            encoding = self._detect_encoding(file_path)

            with open(file_path, 'r', encoding=encoding) as f:
                html_content = f.read()

            if HAS_BS4:
                soup = BeautifulSoup(html_content, 'html.parser')
                # 移除脚本和样式
                for script in soup(["script", "style"]):
                    script.decompose()
                content = soup.get_text(separator='\n', strip=True)
            else:
                # 简单的 HTML 标签移除
                import re
                content = re.sub(r'<[^>]+>', '\n', html_content)

            return {
                "success": True,
                "content": content,
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"读取 HTML 失败: {str(e)}"
            }

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取 HTML 元数据"""
        try:
            encoding = self._detect_encoding(file_path)

            with open(file_path, 'r', encoding=encoding) as f:
                html_content = f.read()

            file_stat = os.stat(file_path)
            metadata = {
                "format": "HTML",
                "encoding": encoding,
                "file_size": file_stat.st_size,
                "created": file_stat.st_ctime,
                "modified": file_stat.st_mtime,
            }

            if HAS_BS4:
                soup = BeautifulSoup(html_content, 'html.parser')
                if soup.title:
                    metadata["title"] = soup.title.string
                metadata["link_count"] = len(soup.find_all('a'))
                metadata["image_count"] = len(soup.find_all('img'))

            return metadata
        except Exception as e:
            return {"error": f"获取元数据失败: {str(e)}"}

    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        if HAS_CHARDET:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                return result['encoding'] or DEFAULT_ENCODING
        return DEFAULT_ENCODING


class TextProcessor(BaseDocumentProcessor):
    """纯文本文档处理器"""

    def __init__(self):
        self.format_name = "text"

    def read(self, file_path: str) -> Dict[str, Any]:
        """读取文本内容"""
        try:
            encoding = self._detect_encoding(file_path)

            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()

            lines = content.split('\n')

            return {
                "success": True,
                "content": content,
                "line_count": len(lines),
                "char_count": len(content),
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"读取文本失败: {str(e)}"
            }

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取文本元数据"""
        try:
            encoding = self._detect_encoding(file_path)

            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()

            file_stat = os.stat(file_path)

            # 统计信息
            words = content.split()
            lines = content.split('\n')

            metadata = {
                "format": "Text",
                "line_count": len(lines),
                "word_count": len(words),
                "char_count": len(content),
                "encoding": encoding,
                "file_size": file_stat.st_size,
                "created": file_stat.st_ctime,
                "modified": file_stat.st_mtime,
            }

            return metadata
        except Exception as e:
            return {"error": f"获取元数据失败: {str(e)}"}

    def _detect_encoding(self, file_path: str) -> str:
        """检测文件编码"""
        if HAS_CHARDET:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                return result['encoding'] or DEFAULT_ENCODING
        return DEFAULT_ENCODING


# ============================================================
# 处理器工厂
# ============================================================

def get_processor(file_path: str) -> BaseDocumentProcessor:
    """根据文件类型获取处理器"""
    format_type = get_format_type(file_path)

    processors = {
        "pdf": PDFProcessor(),
        "word": WordProcessor(),
        "markdown": MarkdownProcessor(),
        "html": HTMLProcessor(),
        "text": TextProcessor(),
    }

    return processors.get(format_type, TextProcessor())


# ============================================================
# MCP 服务器工具处理函数
# ============================================================

async def handle_read_document(file_path: str) -> List[TextContent]:
    """读取文档"""
    try:
        processor = get_processor(file_path)

        # 验证文件
        valid, error_msg = processor.validate_file(file_path)
        if not valid:
            return [TextContent(type="text", text=f"文件验证失败: {error_msg}")]

        # 读取文档
        result = processor.read(file_path)

        if not result.get("success"):
            return [TextContent(type="text", text=result.get("error", "读取失败"))]

        content = result["content"]
        return [TextContent(type="text", text=content)]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


async def handle_extract_metadata(file_path: str) -> List[TextContent]:
    """提取元数据"""
    try:
        processor = get_processor(file_path)

        # 验证文件
        valid, error_msg = processor.validate_file(file_path)
        if not valid:
            return [TextContent(type="text", text=f"文件验证失败: {error_msg}")]

        # 获取元数据
        metadata = processor.get_metadata(file_path)

        if "error" in metadata:
            return [TextContent(type="text", text=metadata["error"])]

        # 格式化输出
        import json
        return [
            TextContent(type="text", text=f"元数据提取成功: {file_path}"),
            TextContent(type="text", text=json.dumps(metadata, indent=2, ensure_ascii=False))
        ]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


async def handle_search_in_document(file_path: str, query: str, case_sensitive: bool = False, context_lines: int = 2) -> List[TextContent]:
    """文档搜索"""
    try:
        processor = get_processor(file_path)

        # 读取文档
        result = processor.read(file_path)
        if not result.get("success"):
            return [TextContent(type="text", text=result.get("error", "读取失败"))]

        content = result["content"]
        lines = content.split('\n')

        # 搜索
        matches = []
        search_query = query if case_sensitive else query.lower()

        for i, line in enumerate(lines):
            search_line = line if case_sensitive else line.lower()
            if search_query in search_line:
                # 获取上下文
                start = max(0, i - context_lines)
                end = min(len(lines), i + context_lines + 1)
                context = lines[start:end]

                matches.append({
                    "line_number": i + 1,
                    "line": line,
                    "context": context
                })

        import json
        return [
            TextContent(type="text", text=f"在 {file_path} 中找到 {len(matches)} 处匹配"),
            TextContent(type="text", text=json.dumps({"matches": matches, "query": query}, indent=2, ensure_ascii=False))
        ]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


async def handle_batch_process(file_paths: List[str], operation: str) -> List[TextContent]:
    """批量处理"""
    try:
        if len(file_paths) > MAX_BATCH_SIZE:
            return [TextContent(type="text", text=f"文件数量超过限制 ({MAX_BATCH_SIZE})")]

        results = []

        for file_path in file_paths:
            processor = get_processor(file_path)

            if operation == "read":
                result = processor.read(file_path)
            elif operation == "metadata":
                result = processor.get_metadata(file_path)
            else:
                result = {"error": f"不支持的操作: {operation}"}

            results.append({
                "file_path": file_path,
                "result": result
            })

        import json
        return [
            TextContent(type="text", text=f"批量处理完成: {len(file_paths)} 个文件"),
            TextContent(type="text", text=json.dumps({"results": results}, indent=2, ensure_ascii=False))
        ]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


async def handle_get_supported_formats() -> List[TextContent]:
    """获取支持格式"""
    import json
    return [
        TextContent(type="text", text="支持的文档格式:"),
        TextContent(type="text", text=json.dumps(SUPPORTED_FORMATS, indent=2, ensure_ascii=False))
    ]


async def handle_validate_document(file_path: str) -> List[TextContent]:
    """验证文档"""
    try:
        processor = get_processor(file_path)
        valid, error_msg = processor.validate_file(file_path)

        if not valid:
            return [TextContent(type="text", text=f"验证失败: {error_msg}")]

        # 尝试读取文件
        result = processor.read(file_path)

        if result.get("success"):
            return [TextContent(type="text", text=f"文档验证通过: {file_path}\n格式: {processor.format_name}")]
        else:
            return [TextContent(type="text", text=f"验证失败: {result.get('error')}")]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


async def handle_convert_to_markdown(file_path: str, output_path: Optional[str] = None) -> List[TextContent]:
    """转换为 Markdown"""
    try:
        processor = get_processor(file_path)

        # 读取文档
        result = processor.read(file_path)
        if not result.get("success"):
            return [TextContent(type="text", text=result.get("error", "读取失败"))]

        content = result["content"]

        # 如果已经是 Markdown，直接返回
        if isinstance(processor, MarkdownProcessor):
            return [TextContent(type="text", text="文档已经是 Markdown 格式")]

        # 生成输出路径
        if not output_path:
            base_path = Path(file_path).stem
            output_path = str(Path(file_path).parent / f"{base_path}.md")

        # 写入 Markdown 文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return [TextContent(type="text", text=f"转换成功: {output_path}")]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


async def handle_get_document_stats(file_path: str) -> List[TextContent]:
    """获取文档统计"""
    try:
        processor = get_processor(file_path)

        # 读取文档
        result = processor.read(file_path)
        if not result.get("success"):
            return [TextContent(type="text", text=result.get("error", "读取失败"))]

        content = result["content"]

        # 统计信息
        stats = {
            "char_count": len(content),
            "char_count_no_spaces": len(content.replace(' ', '').replace('\n', '').replace('\t', '')),
            "line_count": len(content.split('\n')),
            "paragraph_count": len([p for p in content.split('\n\n') if p.strip()]),
        }

        # 估算字数（中文+英文）
        import re
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', content))
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', content))
        stats["word_count"] = chinese_chars + english_words
        stats["chinese_chars"] = chinese_chars
        stats["english_words"] = english_words

        # 估算阅读时间（中文400字/分钟，英文200词/分钟）
        read_time = (chinese_chars / 400) + (english_words / 200)
        stats["estimated_read_time_minutes"] = round(read_time, 2)

        # 获取文件大小
        file_size = os.path.getsize(file_path)
        stats["file_size_bytes"] = file_size
        stats["file_size_kb"] = round(file_size / 1024, 2)

        import json
        return [
            TextContent(type="text", text=f"文档统计: {file_path}"),
            TextContent(type="text", text=json.dumps(stats, indent=2, ensure_ascii=False))
        ]

    except Exception as e:
        return [TextContent(type="text", text=f"错误: {str(e)}")]


# ============================================================
# 创建并运行 MCP 服务器
# ============================================================

# 创建 MCP Server 实例
server = Server("doc_processor")

# 定义工具列表
TOOLS = [
    Tool(
        name="read_document",
        description="读取文档内容，支持 PDF、Word、Markdown、HTML、TXT 格式",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "文档文件路径"}
            },
            "required": ["file_path"]
        }
    ),
    Tool(
        name="extract_metadata",
        description="提取文档元数据（作者、标题、创建时间、页数等）",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "文档文件路径"}
            },
            "required": ["file_path"]
        }
    ),
    Tool(
        name="search_in_document",
        description="在文档中搜索关键词或正则表达式",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "文档文件路径"},
                "query": {"type": "string", "description": "搜索关键词"},
                "case_sensitive": {"type": "boolean", "description": "是否区分大小写", "default": False},
                "context_lines": {"type": "integer", "description": "上下文行数", "default": 2}
            },
            "required": ["file_path", "query"]
        }
    ),
    Tool(
        name="batch_process",
        description="批量处理多个文档",
        inputSchema={
            "type": "object",
            "properties": {
                "file_paths": {"type": "array", "items": {"type": "string"}, "description": "文件路径列表"},
                "operation": {"type": "string", "description": "操作类型: read/metadata"}
            },
            "required": ["file_paths", "operation"]
        }
    ),
    Tool(
        name="get_supported_formats",
        description="获取支持的文档格式列表",
        inputSchema={
            "type": "object",
            "properties": {},
            "required": []
        }
    ),
    Tool(
        name="validate_document",
        description="验证文档的完整性和格式",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "文档文件路径"}
            },
            "required": ["file_path"]
        }
    ),
    Tool(
        name="convert_to_markdown",
        description="将文档转换为 Markdown 格式",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "源文件路径"},
                "output_path": {"type": "string", "description": "输出文件路径（可选）"}
            },
            "required": ["file_path"]
        }
    ),
    Tool(
        name="get_document_stats",
        description="获取文档统计信息（字数、段落数等）",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "文档文件路径"}
            },
            "required": ["file_path"]
        }
    ),
]


@server.list_tools()
async def list_tools() -> List[Tool]:
    """列出可用工具"""
    return TOOLS


@server.call_tool()
async def call_tool(name: str, arguments: Any) -> List[TextContent]:
    """调用工具"""
    if name == "read_document":
        return await handle_read_document(**arguments)
    elif name == "extract_metadata":
        return await handle_extract_metadata(**arguments)
    elif name == "search_in_document":
        return await handle_search_in_document(**arguments)
    elif name == "batch_process":
        return await handle_batch_process(**arguments)
    elif name == "get_supported_formats":
        return await handle_get_supported_formats()
    elif name == "validate_document":
        return await handle_validate_document(**arguments)
    elif name == "convert_to_markdown":
        return await handle_convert_to_markdown(**arguments)
    elif name == "get_document_stats":
        return await handle_get_document_stats(**arguments)
    else:
        return [TextContent(type="text", text=f"未知工具: {name}")]


# ============================================================
# 主程序
# ============================================================

if __name__ == "__main__":
    # 检查命令行参数
    if len(sys.argv) > 1 and sys.argv[1] == "--test":
        # 测试模式：打印信息
        print("=" * 60)
        print("文档处理 MCP 服务器 - 测试模式")
        print("=" * 60)
        print()

        print("依赖检查:")
        print(f"  PyPDF2: {'OK' if HAS_PYPDF2 else 'MISSING (pip install PyPDF2)'}")
        print(f"  python-docx: {'OK' if HAS_DOCX else 'MISSING (pip install python-docx)'}")
        print(f"  markdown: {'OK' if HAS_MARKDOWN else 'MISSING (pip install markdown)'}")
        print(f"  beautifulsoup4: {'OK' if HAS_BS4 else 'MISSING (pip install beautifulsoup4)'}")
        print(f"  chardet: {'OK' if HAS_CHARDET else 'MISSING (pip install chardet)'}")
        print()

        print("已创建 MCP 服务器: doc_processor")
        print("支持的工具:")
        for tool in TOOLS:
            print(f"  - {tool.name}")
        print()
    else:
        # MCP 模式：启动 stdio 服务器
        async def main():
            async with stdio_server() as (read_stream, write_stream):
                await server.run(
                    read_stream,
                    write_stream,
                    server.create_initialization_options()
                )

        asyncio.run(main())
